import requests
import time
import os
from urllib.parse import quote
from plyer import notification
from datetime import datetime
from functools import lru_cache
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

### Kullanıcı ayarları
min_discount = 35
min_price = 50
min_otherprice = 0
max_price = 3000
max_float = 0.010
min_sticker = 500

# Kalıcı veri depolama için dosya adı
seen_ids_file = "seen_products.txt"

# Görülen ürün ID'lerini yükle
def load_seen_ids():
    if not os.path.exists(seen_ids_file):
        return set()
    with open(seen_ids_file, "r") as f:
        return set(line.strip() for line in f)

# Görülen ürün ID'lerini kaydet
def save_seen_ids(seen_ids):
    with open(seen_ids_file, "w") as f:
        for pid in seen_ids:
            f.write(f"{pid}\n")

seen_products_docx = load_seen_ids()

# Word belgesini başlat ve eski ID'leri oku
def init_doc():
    try:
        doc = Document("urunler.docx")
        # Eski ID'leri oku
        for para in doc.paragraphs:
            if "Ürün ID:" in para.text:
                pid = para.text.split("Ürün ID: ")[1].strip()
                seen_products_docx.add(pid)
        return doc
    except:
        return Document()

doc = init_doc()

@lru_cache(maxsize=100)
def fetch_sticker_price(sticker_name):
    encoded_name = quote(sticker_name)
    urll = f"https://gw.bynogame.com/steam-products/v2/products?page=1&limit=36&sort=SalesCountLastSevenDays:-1&filters=Category:Sticker;Keywords:{encoded_name};AppId:730"
    
    try:
        response = requests.get(urll)
        response.raise_for_status()
        dataa = response.json()
        
        if "data" in dataa and "result" in dataa["data"]:
            for result in dataa["data"]["result"]:
                if result.get("marketHashName", "Bilinmeyen sticker") == sticker_name:
                    return result.get("priceTRY", 0)
        
        return 0

    except Exception as e:
        print(f"Sticker fiyatını çekerken hata oluştu: {e}")
        return 0

def notify_product(title, product_name, float_value, product_price, discount_percentage, additional_info=""):
    current_time = datetime.now().strftime("%H:%M:%S")
    
    # Başlık ve mesajı kısalt
    def truncate(text, max_length):
        return text[:max_length-3] + "..." if len(text) > max_length else text
    
    title = truncate(title, 256)
    message = truncate(
        f"{product_name}\nFloat: {float_value}\nFiyat: {product_price:.2f}TL\n%{discount_percentage} indirim\n{additional_info}", 
        256
    )
    
    notification.notify(
        title=title,
        message=message
    )
    print(f"{current_time} - {title}\n{message}\n")

def add_to_docx(product_name, float_value, product_price, discount_percentage, product_ID, additional_info="", title=""):
    if product_ID not in seen_products_docx:
        seen_products_docx.add(product_ID)
        save_seen_ids(seen_products_docx)  # Kalıcı depolama
        
        # Başlık ekle
        heading = doc.add_heading(level=1)
        run = heading.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        
        # İçerik ekle
        content = [
            f"Ürün Adı: {product_name}",
            f"Float Değeri: {float_value}",
            f"Fiyat: {product_price:.2f} TL",
            f"İndirim Oranı: %{discount_percentage}",
            f"Ek Bilgi: {additional_info}" if additional_info else "",
            f"Ürün ID: {product_ID}",
            "-"*40
        ]
        
        for line in content:
            if line:  # Boş satırları atla
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(10)
                p.paragraph_format.line_spacing = 1.0
        
        doc.save("urunler.docx")

def check_for_new_discount_products(products):
    for product in products:
        product_ID = product.get('listingNo', '')
        if product_ID in seen_products_docx:
            continue

        product_name = product.get('name', 'Bilinmeyen Ürün')
        discount_percentage = product.get('discountRate', 0)
        product_price = product.get('price', 0)
        float_value = float(product.get('info', {}).get('float', 1))

        if min_price <= product_price <= max_price and discount_percentage >= min_discount:
            notify_product("İNDİRİMLİ ÜRÜN", product_name, float_value, product_price, discount_percentage)
            add_to_docx(product_name, float_value, product_price, discount_percentage, product_ID, title="İNDİRİMLİ ÜRÜN")

def check_for_new_float_products(products):
    for product in products:
        product_ID = product.get('listingNo', '')
        if product_ID in seen_products_docx:
            continue

        product_name = product.get('name', 'Bilinmeyen Ürün')
        discount_percentage = product.get('discountRate', 0)
        product_price = product.get('price', 0)
        float_value = float(product.get('info', {}).get('float', 1))

        if min_otherprice <= product_price <= max_price and -1 < float_value < max_float:
            notify_product("DÜŞÜK FLOAT", product_name, float_value, product_price, discount_percentage)
            add_to_docx(product_name, float_value, product_price, discount_percentage, product_ID, title="DÜŞÜK FLOAT")

def check_for_valuable_sticker_products(products):
    for product in products:
        product_ID = product.get('listingNo', '')
        if product_ID in seen_products_docx:
            continue

        product_name = product.get('name', 'Bilinmeyen Ürün')
        product_price = product.get('price', 0)
        discount_percentage = product.get('discountRate', 0)
        float_value = float(product.get('info', {}).get('float', 1))
        stickers = product.get('info', {}).get('stickerNames', 'Sticker Yok')

        if stickers != 'Sticker Yok' and 'souvenir' not in product_name.lower():
            sticker_list = stickers.split(', Sticker | ')
            total_sticker_value = sum(fetch_sticker_price(sticker.strip()) for sticker in sticker_list)
            
            if total_sticker_value >= min_sticker and min_otherprice <= product_price <= max_price:
                notify_product("DEĞERLİ STİCKER", product_name, float_value, product_price, discount_percentage, f"Toplam: {total_sticker_value:.2f}TL")
                add_to_docx(
                    product_name, 
                    float_value, 
                    product_price, 
                    discount_percentage, 
                    product_ID, 
                    f"Stickerlar: {', '.join(sticker_list)}\nToplam Değer: {total_sticker_value:.2f}TL", 
                    title="DEĞERLİ STİCKER"
                )

while True:
    try:
        url = "https://listing.bynogame.net/api/listings/cs2?page=0&limit=1000"
        response = requests.get(url)
        products = response.json().get('payload', [])
        
        check_for_new_discount_products(products)
        check_for_new_float_products(products)
        check_for_valuable_sticker_products(products)
        
    except Exception as e:
        print(f"Hata: {e}")
    
    time.sleep(10)